from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.vectorstores import Chroma
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import docx
import os
import re

DOC_REPO_DIRECTORY = "documents/"
HF_EMBEDDINGS = 'BAAI/bge-large-en-v1.5'
CHROMA_LOCAL = "chroma_db"

embeddings = HuggingFaceInstructEmbeddings(model_name=HF_EMBEDDINGS)
vector_store = Chroma(persist_directory=CHROMA_LOCAL, embedding_function=embeddings)

def preprocess_docs():    
    doc_paths = []

    w = os.walk(DOC_REPO_DIRECTORY)
    for (path, name, file_name) in w:
        for file in file_name:
            try:
                doc = docx.Document(path + "/" + file)
                doc_paths.append(path + "/" + file)
            except Exception:
                pass
    
    for doc_path in doc_paths:
        ingest_doc(doc_path)

def get_link_from_text(text):
    pattern = r'(https://\S+)'

    match = re.search(pattern, text)

    if match:
        link = match.group(1)
        return link
    else:
        return "None"
    
def ingest_doc(doc_path):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200        
    )
    doc = docx.Document(doc_path)
    link = get_link_from_text(doc.paragraphs[0].text)
    content = ""
    for p in doc.paragraphs[1:]:
        content += p.text

    chunks = text_splitter.split_text(text=content)
    for i in range(0, len(chunks)):                
        new_doc = Document(
            page_content=chunks[i],
            metadata={
                "source": link,                        
            }
        )
        vector_store.add_documents([new_doc], ids=[doc_path + "_chunk" + str(i)])
            
if __name__ == '__main__':
    preprocess_docs()